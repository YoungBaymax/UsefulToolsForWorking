#!/usr/bin/env python
# -*- encoding: utf-8 -*-
'''
@File    :   docxWriter.py    
@Contact :   raogx.vip@hotmail.com
@License :   (C)Copyright 2017-2018, Liugroup-NLPR-CASIA

@Modify Time      @Author    @Version    @Desciption
------------      -------    --------    -----------
2021/2/23 14:31   gxrao      1.0         None
'''

# import lib
import docx.package
from docx import Document
from docx.shared import Inches
from docx.shared import Cm
import os
import docx2pdf

report_doc = Document()

# title
mainTitle = report_doc.add_heading('上海市第十人民医院心血管内科', 0)
secondTitle = report_doc.add_heading('心脏压力负荷报告', level = 0)
mainTitle.add_run()
secondTitle.add_run('font')

p = report_doc.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

report_doc.add_heading('Heading, level 1', level=1)
report_doc.add_paragraph('Intense quote', style= 'Intense Quote')
report_doc.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
report_doc.add_paragraph(
    'first item in ordered list', style='List Number'
)
report_doc.add_picture('lombplot.png', width=Cm(14.8), height = Cm(5.86))

records = (
    (3, '101', '1号'),
    (7, '422', '二号'),
    (4, '631', 'Spam, spam, eggs, and spam')
)
table = report_doc.add_table(rows=1, cols=3, style= "Light Shading Accent 1")
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '正'
hdr_cells[1].text = '价'
hdr_cells[2].text = '它'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc
report_doc.add_page_break()
wkdir = os.curdir
file_name = os.path.join(wkdir, "testDemo.docx")
if os.path.exists(file_name):
    print(file_name, "has been removed.")
    os.remove(file_name)
report_doc.save(file_name)
docx2pdf.word2pdf("./testDemo.docx")